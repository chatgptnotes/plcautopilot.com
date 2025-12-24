"""
Create Word Document for 3-Pump 3-Tank Backup System Documentation
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os


def create_documentation():
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('3-Pump 3-Tank Backup System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Program Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.bold = True

    # Overview Section
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        'This PLC program controls a water distribution system with 3 pumps and 3 tanks. '
        'Each pump is dedicated to filling one tank, but the system includes backup logic '
        'so that if a pump fails, a neighboring pump can take over and fill the affected '
        'tank through interconnecting valves.'
    )

    # Add specs table
    specs_table = doc.add_table(rows=3, cols=2)
    specs_table.style = 'Table Grid'
    specs_data = [
        ('Controller', 'Schneider Electric TM221CE24T (Modicon M221)'),
        ('Software', 'EcoStruxure Machine Expert - Basic'),
        ('File', 'Pump_Tank_TM221.smbp')
    ]
    for i, (label, value) in enumerate(specs_data):
        specs_table.rows[i].cells[0].text = label
        specs_table.rows[i].cells[1].text = value
        specs_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    # System Architecture
    doc.add_heading('2. System Architecture', level=1)

    arch_text = """
                        WATER SOURCE
                             |
             +---------------+---------------+
             |               |               |
         [PUMP 1]        [PUMP 2]        [PUMP 3]
             |               |               |
             v               v               v
         +-------+       +-------+       +-------+
         |       |       |       |       |       |
         |TANK 1 |<----->|TANK 2 |<----->|TANK 3 |
         |       |VALVE12|       |VALVE23|       |
         +-------+       +-------+       +-------+
    """

    arch_para = doc.add_paragraph(arch_text)
    arch_para.runs[0].font.name = 'Courier New'
    arch_para.runs[0].font.size = Pt(9)

    doc.add_heading('Normal Operation', level=2)
    doc.add_paragraph('Pump 1 fills Tank 1', style='List Bullet')
    doc.add_paragraph('Pump 2 fills Tank 2', style='List Bullet')
    doc.add_paragraph('Pump 3 fills Tank 3', style='List Bullet')
    doc.add_paragraph('Valves between tanks remain CLOSED', style='List Bullet')

    doc.add_heading('Backup Operation', level=2)
    doc.add_paragraph('If Pump 1 fails: Pump 2 takes over Tank 1, Valve 1-2 OPENS', style='List Bullet')
    doc.add_paragraph('If Pump 2 fails: Pump 3 takes over Tank 2, Valve 2-3 OPENS', style='List Bullet')

    # I/O Assignment
    doc.add_heading('3. I/O Assignment', level=1)

    # Digital Inputs Table
    doc.add_heading('Digital Inputs (12 of 14 used)', level=2)

    input_table = doc.add_table(rows=13, cols=4)
    input_table.style = 'Table Grid'

    input_headers = ['Address', 'Symbol', 'Description', 'Type']
    for i, header in enumerate(input_headers):
        input_table.rows[0].cells[i].text = header
        input_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    input_data = [
        ('%I0.0', 'START_BTN', 'System Start Button', 'NO'),
        ('%I0.1', 'STOP_BTN', 'System Stop Button', 'NC'),
        ('%I0.2', 'PUMP1_SPEED_OK', 'Pump 1 Zero Speed Sensor', 'NC'),
        ('%I0.3', 'PUMP2_SPEED_OK', 'Pump 2 Zero Speed Sensor', 'NC'),
        ('%I0.4', 'PUMP3_SPEED_OK', 'Pump 3 Zero Speed Sensor', 'NC'),
        ('%I0.5', 'TANK1_LOW', 'Tank 1 Low Level Switch', 'NO'),
        ('%I0.6', 'TANK1_HIGH', 'Tank 1 High Level Switch', 'NO'),
        ('%I0.7', 'TANK2_LOW', 'Tank 2 Low Level Switch', 'NO'),
        ('%I0.8', 'TANK2_HIGH', 'Tank 2 High Level Switch', 'NO'),
        ('%I0.9', 'TANK3_LOW', 'Tank 3 Low Level Switch', 'NO'),
        ('%I0.10', 'TANK3_HIGH', 'Tank 3 High Level Switch', 'NO'),
        ('%I0.11', 'FAULT_RESET', 'Fault Reset Button', 'NO'),
    ]

    for i, row_data in enumerate(input_data):
        for j, cell_data in enumerate(row_data):
            input_table.rows[i+1].cells[j].text = cell_data

    doc.add_paragraph()
    doc.add_paragraph(
        'Note: Zero speed sensors are wired NC (Normally Closed). When the motor is running '
        'at speed, the contact is CLOSED (signal = 1). When motor stops or fails, contact OPENS (signal = 0).'
    ).runs[0].font.italic = True

    # Digital Outputs Table
    doc.add_heading('Digital Outputs (10 of 10 used)', level=2)

    output_table = doc.add_table(rows=11, cols=4)
    output_table.style = 'Table Grid'

    output_headers = ['Address', 'Symbol', 'Description', 'Load Type']
    for i, header in enumerate(output_headers):
        output_table.rows[0].cells[i].text = header
        output_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    output_data = [
        ('%Q0.0', 'PUMP1_RUN', 'Pump 1 Motor Contactor', 'Relay/24VDC'),
        ('%Q0.1', 'PUMP2_RUN', 'Pump 2 Motor Contactor', 'Relay/24VDC'),
        ('%Q0.2', 'PUMP3_RUN', 'Pump 3 Motor Contactor', 'Relay/24VDC'),
        ('%Q0.3', 'VALVE_12', 'Valve between Tank 1-2', 'Solenoid'),
        ('%Q0.4', 'VALVE_23', 'Valve between Tank 2-3', 'Solenoid'),
        ('%Q0.5', 'PUMP1_FAULT_IND', 'Pump 1 Fault Indicator', 'Lamp'),
        ('%Q0.6', 'PUMP2_FAULT_IND', 'Pump 2 Fault Indicator', 'Lamp'),
        ('%Q0.7', 'PUMP3_FAULT_IND', 'Pump 3 Fault Indicator', 'Lamp'),
        ('%Q0.8', 'SYSTEM_RUN_IND', 'System Running Indicator', 'Lamp'),
        ('%Q0.9', 'ALARM_OUTPUT', 'Alarm Horn/Buzzer', 'Horn'),
    ]

    for i, row_data in enumerate(output_data):
        for j, cell_data in enumerate(row_data):
            output_table.rows[i+1].cells[j].text = cell_data

    # Memory Bits Table
    doc.add_heading('Memory Bits', level=2)

    memory_table = doc.add_table(rows=10, cols=3)
    memory_table.style = 'Table Grid'

    memory_headers = ['Address', 'Symbol', 'Description']
    for i, header in enumerate(memory_headers):
        memory_table.rows[0].cells[i].text = header
        memory_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    memory_data = [
        ('%M0', 'SYSTEM_RUN', 'System Running Status'),
        ('%M1', 'PUMP1_CMD', 'Pump 1 Fill Command'),
        ('%M2', 'PUMP2_CMD', 'Pump 2 Fill Command'),
        ('%M3', 'PUMP3_CMD', 'Pump 3 Fill Command'),
        ('%M10', 'PUMP1_FAULT', 'Pump 1 Fault Latch'),
        ('%M11', 'PUMP2_FAULT', 'Pump 2 Fault Latch'),
        ('%M12', 'PUMP3_FAULT', 'Pump 3 Fault Latch'),
        ('%M20', 'TANK1_BACKUP', 'Tank 1 Needs Backup (Pump 2)'),
        ('%M21', 'TANK2_BACKUP', 'Tank 2 Needs Backup (Pump 3)'),
    ]

    for i, row_data in enumerate(memory_data):
        for j, cell_data in enumerate(row_data):
            memory_table.rows[i+1].cells[j].text = cell_data

    # Timers Table
    doc.add_heading('Timers', level=2)

    timer_table = doc.add_table(rows=4, cols=6)
    timer_table.style = 'Table Grid'

    timer_headers = ['Address', 'Symbol', 'Type', 'Time Base', 'Preset', 'Description']
    for i, header in enumerate(timer_headers):
        timer_table.rows[0].cells[i].text = header
        timer_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    timer_data = [
        ('%TM0', 'PUMP1_DELAY', 'TON', '1 second', '2', 'Pump 1 startup delay'),
        ('%TM1', 'PUMP2_DELAY', 'TON', '1 second', '2', 'Pump 2 startup delay'),
        ('%TM2', 'PUMP3_DELAY', 'TON', '1 second', '2', 'Pump 3 startup delay'),
    ]

    for i, row_data in enumerate(timer_data):
        for j, cell_data in enumerate(row_data):
            timer_table.rows[i+1].cells[j].text = cell_data

    # Program Logic
    doc.add_heading('4. Program Logic (19 Rungs)', level=1)

    # Rung 1
    doc.add_heading('Rung 1: System Start/Stop Control', level=2)

    rung1_ladder = """
     START_BTN        STOP_BTN(NC)
--+----[/I0.0]----+----[/I0.1]-------------------(M0)---
  |               |                              SYSTEM_RUN
  +----[/M0]------+
      SEAL-IN
"""
    rung1_para = doc.add_paragraph(rung1_ladder)
    rung1_para.runs[0].font.name = 'Courier New'
    rung1_para.runs[0].font.size = Pt(9)

    doc.add_paragraph('Logic:', style='List Bullet')
    doc.add_paragraph('Press START button to energize SYSTEM_RUN', style='List Bullet')
    doc.add_paragraph('SYSTEM_RUN seals in (latches) through parallel contact', style='List Bullet')
    doc.add_paragraph('Press STOP button (NC) to de-energize and stop system', style='List Bullet')

    # Rungs 2-4
    doc.add_heading('Rungs 2-4: Pump Command Logic', level=2)
    doc.add_paragraph('Each pump command follows this pattern:')

    rung24_ladder = """
   SYSTEM_RUN    TANK_LOW     TANK_HIGH(NC)  PUMP_FAULT(NC)
----[/M0]-------[/I0.x]-------[/I0.y]--------[/M1x]-------(Mx)---
                                                          PUMP_CMD
"""
    rung24_para = doc.add_paragraph(rung24_ladder)
    rung24_para.runs[0].font.name = 'Courier New'
    rung24_para.runs[0].font.size = Pt(9)

    doc.add_paragraph('Conditions for pump to command ON:')
    doc.add_paragraph('System must be running (%M0 = 1)', style='List Number')
    doc.add_paragraph('Tank low level switch active (tank needs water)', style='List Number')
    doc.add_paragraph('Tank high level switch NOT active (tank not full)', style='List Number')
    doc.add_paragraph('No fault on this pump', style='List Number')

    # Rungs 6, 8, 10 - Fault Detection
    doc.add_heading('Rungs 6, 8, 10: Zero Speed Fault Detection', level=2)

    fault_ladder = """
   TIMER.Q       SPEED_OK(NC)   FAULT_RESET(NC)
--+---[/TMx.Q]-----[/I0.x]---+----[/I0.11]------------(M1x)---
  |                          |                        PUMP_FAULT
  +--------[/M1x]------------+
           LATCH
"""
    fault_para = doc.add_paragraph(fault_ladder)
    fault_para.runs[0].font.name = 'Courier New'
    fault_para.runs[0].font.size = Pt(9)

    doc.add_paragraph('Fault Detection Logic:')
    doc.add_paragraph('After timer expires (pump has had 2 seconds to reach speed)', style='List Number')
    doc.add_paragraph('If SPEED_OK signal is FALSE (motor not spinning)', style='List Number')
    doc.add_paragraph('Then FAULT is SET and LATCHED', style='List Number')
    doc.add_paragraph('Fault remains latched until FAULT_RESET button is pressed', style='List Number')

    doc.add_paragraph(
        'Why 2-second delay? Motors need time to accelerate. The delay prevents false '
        'fault detection during startup.'
    ).runs[0].font.italic = True

    # Rungs 11-12 - Backup Request
    doc.add_heading('Rungs 11-12: Backup Request Logic', level=2)

    doc.add_paragraph('Tank 1 Backup (Rung 11):')
    backup1_ladder = """
   PUMP1_FAULT   TANK1_LOW    PUMP2_FAULT(NC)
----[/M10]-------[/I0.5]-------[/M11]-------------------(M20)---
                                                      TANK1_BACKUP
"""
    backup1_para = doc.add_paragraph(backup1_ladder)
    backup1_para.runs[0].font.name = 'Courier New'
    backup1_para.runs[0].font.size = Pt(9)

    doc.add_paragraph(
        'If Pump 1 has faulted AND Tank 1 needs water AND Pump 2 is healthy, request backup.'
    )

    doc.add_paragraph('Tank 2 Backup (Rung 12):')
    backup2_ladder = """
   PUMP2_FAULT   TANK2_LOW    PUMP3_FAULT(NC)
----[/M11]-------[/I0.7]-------[/M12]-------------------(M21)---
                                                      TANK2_BACKUP
"""
    backup2_para = doc.add_paragraph(backup2_ladder)
    backup2_para.runs[0].font.name = 'Courier New'
    backup2_para.runs[0].font.size = Pt(9)

    doc.add_paragraph(
        'If Pump 2 has faulted AND Tank 2 needs water AND Pump 3 is healthy, request backup.'
    )

    # Rungs 13-14 - Valve Control
    doc.add_heading('Rungs 13-14: Valve Control', level=2)
    doc.add_paragraph(
        'Valves open automatically when backup mode is active, allowing the backup pump '
        'to fill the affected tank through the interconnection.'
    )
    doc.add_paragraph('TANK1_BACKUP (%M20) -> VALVE_12 (%Q0.3)', style='List Bullet')
    doc.add_paragraph('TANK2_BACKUP (%M21) -> VALVE_23 (%Q0.4)', style='List Bullet')

    # Rungs 15-19 - Indicators
    doc.add_heading('Rungs 15-19: Indicators and Alarm', level=2)
    doc.add_paragraph('Rung 15: PUMP1_FAULT (%M10) -> PUMP1_FAULT_IND (%Q0.5)', style='List Bullet')
    doc.add_paragraph('Rung 16: PUMP2_FAULT (%M11) -> PUMP2_FAULT_IND (%Q0.6)', style='List Bullet')
    doc.add_paragraph('Rung 17: PUMP3_FAULT (%M12) -> PUMP3_FAULT_IND (%Q0.7)', style='List Bullet')
    doc.add_paragraph('Rung 18: SYSTEM_RUN (%M0) -> SYSTEM_RUN_IND (%Q0.8)', style='List Bullet')
    doc.add_paragraph('Rung 19: Any Fault -> ALARM_OUTPUT (%Q0.9)', style='List Bullet')

    # Operating Procedures
    doc.add_heading('5. Operating Procedures', level=1)

    doc.add_heading('System Startup', level=2)
    doc.add_paragraph('Verify all tanks have correct level sensor readings', style='List Number')
    doc.add_paragraph('Verify all pumps are ready (no mechanical issues)', style='List Number')
    doc.add_paragraph('Verify zero speed sensors are functional', style='List Number')
    doc.add_paragraph('Press START button', style='List Number')
    doc.add_paragraph('System Running indicator illuminates', style='List Number')
    doc.add_paragraph('Pumps will automatically start filling low tanks', style='List Number')

    doc.add_heading('Fault Handling', level=2)
    doc.add_paragraph('When a pump fault occurs:')
    doc.add_paragraph('Fault indicator lamp illuminates', style='List Number')
    doc.add_paragraph('Alarm sounds', style='List Number')
    doc.add_paragraph('Backup pump (if available) takes over', style='List Number')
    doc.add_paragraph('Interconnecting valve opens', style='List Number')

    doc.add_paragraph('To clear a fault:')
    doc.add_paragraph('Fix the mechanical issue with the pump', style='List Number')
    doc.add_paragraph('Verify motor can spin freely', style='List Number')
    doc.add_paragraph('Press FAULT_RESET button', style='List Number')
    doc.add_paragraph('Fault indicator extinguishes', style='List Number')
    doc.add_paragraph('Normal operation resumes', style='List Number')

    doc.add_heading('System Shutdown', level=2)
    doc.add_paragraph('Press STOP button', style='List Number')
    doc.add_paragraph('All pumps stop immediately', style='List Number')
    doc.add_paragraph('All valves close', style='List Number')
    doc.add_paragraph('System Running indicator extinguishes', style='List Number')

    # Troubleshooting
    doc.add_heading('6. Troubleshooting', level=1)

    trouble_table = doc.add_table(rows=7, cols=3)
    trouble_table.style = 'Table Grid'

    trouble_headers = ['Symptom', 'Possible Cause', 'Solution']
    for i, header in enumerate(trouble_headers):
        trouble_table.rows[0].cells[i].text = header
        trouble_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    trouble_data = [
        ("Pump won't start", 'Tank high level active', 'Check level sensors'),
        ("Pump won't start", 'Pump fault latched', 'Reset fault, check motor'),
        ('False fault alarm', 'Zero speed sensor failed', 'Replace sensor'),
        ('False fault alarm', 'Timer too short', 'Increase preset (currently 2s)'),
        ('Backup not working', 'Backup pump also faulted', 'Repair backup pump'),
        ('Valve stuck closed', 'Solenoid failed', 'Check solenoid coil'),
    ]

    for i, row_data in enumerate(trouble_data):
        for j, cell_data in enumerate(row_data):
            trouble_table.rows[i+1].cells[j].text = cell_data

    # Maintenance Schedule
    doc.add_heading('7. Maintenance Schedule', level=1)

    maint_table = doc.add_table(rows=6, cols=2)
    maint_table.style = 'Table Grid'

    maint_headers = ['Task', 'Frequency']
    for i, header in enumerate(maint_headers):
        maint_table.rows[0].cells[i].text = header
        maint_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    maint_data = [
        ('Test START/STOP buttons', 'Weekly'),
        ('Test FAULT_RESET button', 'Weekly'),
        ('Verify level sensor operation', 'Monthly'),
        ('Test zero speed sensors', 'Monthly'),
        ('Full system test (simulate faults)', 'Quarterly'),
    ]

    for i, row_data in enumerate(maint_data):
        for j, cell_data in enumerate(row_data):
            maint_table.rows[i+1].cells[j].text = cell_data

    # Document Info
    doc.add_heading('8. Document Information', level=1)

    info_table = doc.add_table(rows=7, cols=2)
    info_table.style = 'Table Grid'

    info_data = [
        ('Program Version', '1.0'),
        ('Created', 'December 2024'),
        ('Controller', 'TM221CE24T'),
        ('Total Rungs', '19'),
        ('Inputs Used', '12 of 14'),
        ('Outputs Used', '10 of 10'),
        ('Timers Used', '3'),
    ]

    for i, row_data in enumerate(info_data):
        info_table.rows[i].cells[0].text = row_data[0]
        info_table.rows[i].cells[1].text = row_data[1]
        info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

    # Save document
    output_path = os.path.join(
        os.path.expanduser("~"),
        "OneDrive",
        "Documents",
        "Pump_Tank_TM221_Documentation.docx"
    )

    doc.save(output_path)
    print(f"Document created: {output_path}")

    # Also save to project folder
    project_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "Pump_Tank_TM221_Documentation.docx"
    )
    doc.save(project_path)
    print(f"Document created: {project_path}")

    return output_path


if __name__ == "__main__":
    print("Creating Word Document...")
    print("=" * 60)
    create_documentation()
    print("=" * 60)
    print("Done!")
